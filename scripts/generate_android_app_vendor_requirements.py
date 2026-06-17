from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "design"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "GCS-Saker_안드로이드앱_개발요구사항서_v0.1.docx"
PDF_PATH = OUT_DIR / "GCS-Saker_안드로이드앱_개발요구사항서_v0.1.pdf"

DASHBOARD_SCREENSHOT = ASSET_DIR / "gcs_saker_current_dashboard_20260608.png"
CCTV_SCREENSHOT = ASSET_DIR / "gcs_saker_current_cctv_20260608.png"
EVENT_LOG_SCREENSHOT = ASSET_DIR / "gcs_saker_current_event_log_20260608.png"
SETTINGS_SCREENSHOT = ASSET_DIR / "gcs_saker_current_settings_20260608.png"

FONT_NAME = "Apple SD Gothic Neo"
INK = RGBColor(22, 32, 44)
MUTED = RGBColor(82, 96, 112)
BLUE = RGBColor(31, 105, 160)
TEAL = RGBColor(15, 118, 110)
RED = RGBColor(153, 27, 27)
LIGHT_BLUE = "EAF3FA"
LIGHT_TEAL = "E6F4F1"
LIGHT_RED = "FDECEC"
BORDER = "D3DCE6"


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 9.2) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def configure_section(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_document(doc: Document) -> None:
    configure_section(doc.sections[0], landscape=False)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.13
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("GCS-Saker Android App Requirements")
        set_font(run, size=8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("A4AI / Android 앱 개발사 전달용 / v0.1")
        set_font(run, size=8.2, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else TEAL)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.13
    run = paragraph.add_run(text)
    set_font(run, size=10.2, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=9.9, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, color: RGBColor = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_table_borders(table, color="B7C9D8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=10.5, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(body)
    set_font(run, size=9.7, color=INK)


def add_table(doc: Document, header: tuple[str, ...], rows: list[tuple[str, ...]], widths: tuple[float, ...]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table)
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, text, bold=True, color=BLUE, size=9.4)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_margins(cells[idx])
            if idx == 0:
                set_cell_shading(cells[idx], "F8FAFC")
            set_cell_text(cells[idx], text, bold=(idx == 0), size=8.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("GCS-Saker")
    set_font(run, size=28, bold=True, color=BLUE)
    p = doc.add_paragraph()
    run = p.add_run("안드로이드 모바일 앱 개발 요구사항서")
    set_font(run, size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("현 웹 관제 시스템과 연동되는 Android Native 앱 개발 범위")
    set_font(run, size=12.5, color=MUTED)
    add_callout(
        doc,
        "중요: 웹 사이트 개선 요구가 아닙니다",
        "본 문서는 현 GCS-Saker 웹 관제 화면을 참고하여 Android 모바일 앱을 개발하기 위한 요구사항입니다. "
        "웹 대시보드 디자인 개선, 백엔드 신규 구축, 서버 인프라 운영은 본 개발 범위가 아니며, 앱은 A4AI가 제공하는 인증, 스트리밍, 지도, 이벤트, 운영 API와 연동합니다.",
        LIGHT_RED,
        RED,
    )
    add_table(
        doc,
        ("구분", "내용"),
        [
            ("대상 플랫폼", "Android Native 앱. iOS는 현재 범위에서 제외"),
            ("권장 기술", "Kotlin, Jetpack Compose, Android WebRTC native SDK, Gradle, Android Studio"),
            ("핵심 목적", "모바일 카메라 송출, 스트림 수신, 지도/GPS, 이벤트 알림, 운영 상태 확인"),
            ("연동 기준", "GCS-Saker 서버의 Auth, Stream, Signaling, Telemetry, Event, Time API"),
            ("납품 기준", "소스코드, 테스트, 빌드 산출물, 운영 매뉴얼, API 연동 문서, 릴리즈 노트"),
        ],
        (1.55, 4.75),
    )


def add_image_pair_page(
    doc: Document,
    heading: str,
    body: str,
    left: tuple[str, Path],
    right: tuple[str, Path],
) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    add_header_footer(doc)
    add_heading(doc, heading)
    add_paragraph(doc, body)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="EEF2F6")
    for idx, (caption, path) in enumerate((left, right)):
        cell = table.cell(0, idx)
        set_cell_margins(cell, top=90, bottom=90, start=90, end=90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(caption)
        set_font(run, size=9.3, bold=True, color=BLUE)
        if path.exists():
            cell.add_paragraph().add_run().add_picture(str(path), width=Inches(4.45))
        else:
            p = cell.add_paragraph()
            run = p.add_run(f"이미지 누락: {path.name}")
            set_font(run, size=8.5, color=RED)


def add_reference_screens(doc: Document) -> None:
    add_image_pair_page(
        doc,
        "참고 화면 1: 대시보드 / CCTV",
        "Android 앱이 연동해야 하는 현 웹 관제 시스템의 주요 화면입니다. 모바일 앱은 이 화면을 그대로 축소 복제하지 않고, 모바일 사용 맥락에 맞게 재구성합니다.",
        ("메인 대시보드", DASHBOARD_SCREENSHOT),
        ("CCTV", CCTV_SCREENSHOT),
    )
    add_image_pair_page(
        doc,
        "참고 화면 2: 이벤트로그 / 운영설정",
        "이벤트 조회, 운영 상태 확인, 시간 동기화 설정은 모바일 앱에서도 접근 가능한 구조로 재해석합니다.",
        ("이벤트로그", EVENT_LOG_SCREENSHOT),
        ("운영설정", SETTINGS_SCREENSHOT),
    )


def add_requirements_pages(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    add_header_footer(doc)

    add_heading(doc, "1. 앱 사용자와 핵심 시나리오")
    add_table(
        doc,
        ("사용자", "주요 목적", "필수 화면/기능"),
        [
            ("현장 송출자", "휴대폰 카메라/마이크/GPS를 이용해 현장 영상을 서버로 송출", "로그인, 송출 준비, WebRTC signaling 상태, 카메라 송출, GPS 전송, 송출 종료"),
            ("관제 사용자", "모바일에서 주요 스트림과 이벤트를 확인", "스트림 목록, 선택 스트림 재생, 지도/위치, 이벤트 알림, 서버 상태 요약"),
            ("운영 관리자", "장애 알림과 설정 상태를 확인", "푸시 알림, 이벤트 로그 필터, 시간 동기화 상태, 서버 구성 상태"),
        ],
        (1.35, 2.35, 2.65),
    )

    add_heading(doc, "2. 필수 기능 요구사항")
    add_table(
        doc,
        ("영역", "요구사항", "수용 기준"),
        [
            ("인증/인가", "서버 로그인, 토큰 갱신, 로그아웃, 역할/그룹 권한 반영", "만료 토큰은 자동 갱신 또는 재로그인 유도. 토큰은 Android Keystore 기반 보안 저장소 사용"),
            ("카메라 송출", "전/후면 카메라, 마이크, 해상도/비트레이트 선택, 송출 시작/정지", "송출 전 signaling 준비 상태 표시. 송출 실패 시 원인과 재시도 안내"),
            ("WebRTC 수신", "스트림 목록 선택, WHEP/WebRTC 수신, 오디오 재생, 연결 끊김 감지", "끊김/재연결/오프라인 상태가 UI에 명확히 표시"),
            ("STUN/TURN", "서버 API에서 ICE 서버 목록을 받아 사용", "앱에 STUN/TURN 주소를 하드코딩하지 않음. TURN fallback 상태를 진단 화면에 표시"),
            ("GPS/지도", "송출 중 위치 전송, 수신 스트림 위치 표시, 지도 포커스", "권한 거부/위치 실패 상태 처리. 공개망/폐쇄망 지도 전략 분리 가능"),
            ("이벤트/알림", "이벤트 로그 조회, 심각도 필터, 푸시 알림", "Firebase Cloud Messaging 기반 푸시. 폐쇄망은 WebSocket/폴링 알림 fallback 설계 포함"),
            ("운영 설정", "시간 동기화 상태, API 서버 연결 상태, 앱 진단 정보", "사용자에게 필요한 진단만 노출하고 비밀값은 표시하지 않음"),
        ],
        (1.15, 3.05, 2.15),
    )

    doc.add_page_break()

    add_heading(doc, "3. 모바일 UX 요구사항")
    for item in [
        "웹 화면을 단순 축소하지 않고 모바일 우선 구조로 재설계한다.",
        "송출 화면은 한 손 조작이 가능해야 하며, 시작/중지/상태/권한/네트워크 품질이 즉시 보여야 한다.",
        "관제 화면은 선택 스트림, 지도, 이벤트를 탭 또는 하단 내비게이션으로 빠르게 전환할 수 있어야 한다.",
        "장애, 연결 지연, TURN fallback, 송출 끊김, 권한 거부 등은 색상만이 아니라 텍스트와 아이콘으로도 구분한다.",
        "배터리/발열/네트워크 상태가 송출 품질에 영향을 주는 경우 사용자에게 조절 가능한 품질 옵션을 제공한다.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 보안 요구사항")
    add_table(
        doc,
        ("항목", "요구사항", "금지/주의"),
        [
            ("토큰 저장", "Access/refresh token은 Android Keystore 또는 EncryptedSharedPreferences 사용", "일반 SharedPreferences, 로그, 화면, crash report에 토큰 노출 금지"),
            ("통신", "모든 운영 통신은 HTTPS/WSS 사용. 사설 인증서 환경은 설치/검증 절차 문서화", "HTTP 평문 운영 금지"),
            ("권한", "카메라, 마이크, 위치 권한은 목적과 상태를 명확히 안내", "권한을 거부해도 앱이 비정상 종료되면 안 됨"),
            ("로그", "민감 정보 제거 후 앱 진단 로그 제공", "비밀번호, 토큰, TURN credential, 개인 위치 원문 과다 저장 금지"),
            ("그룹 권한", "서버가 내려주는 역할/그룹 권한에 따라 스트림 목록과 이벤트 접근 제한", "앱 내부 우회로 숨겨진 스트림 접근 금지"),
        ],
        (1.1, 3.05, 2.2),
    )

    doc.add_page_break()

    add_heading(doc, "5. 성능과 안정성 기준")
    add_table(
        doc,
        ("지표", "목표", "검증 방식"),
        [
            ("송출 준비", "로그인 후 송출 준비 상태까지 3초 이내 목표", "실기기 Wi-Fi/LTE 환경 반복 측정"),
            ("영상 지연", "LAN WebRTC 기준 500ms 내외 목표, TURN 경유 시 1.5초 이내 목표", "타임코드 화면 촬영 또는 서버/클라이언트 timestamp 비교"),
            ("재연결", "일시적 네트워크 단절 후 자동 재시도와 상태 표시", "Wi-Fi 끊김, LTE 전환, 앱 백그라운드 복귀 테스트"),
            ("발열/배터리", "30분 이상 송출 시 과도한 발열/프레임 드랍 완화", "해상도/비트레이트 단계별 실기기 테스트"),
            ("오디오", "영상과 함께 오디오 수신/송출 지원. 지연 상태 표시", "마이크 권한, 음소거, 네트워크 불안정 테스트"),
        ],
        (1.3, 2.75, 2.3),
    )

    add_heading(doc, "6. 테스트와 납품 기준")
    add_table(
        doc,
        ("구분", "필수 산출물", "검수 기준"),
        [
            ("소스코드", "Android Studio 프로젝트, Gradle 설정, 환경 분리", "빌드 재현 가능. secret은 repository에 포함하지 않음"),
            ("테스트", "Unit, UI, integration, WebRTC signaling smoke test", "CI 또는 로컬 명령으로 통과 결과 확인 가능"),
            ("실기기 검증", "최소 2종 이상 Android 기기에서 송출/수신/푸시/위치 테스트", "테스트 기기, OS 버전, 네트워크 조건 보고"),
            ("문서", "설치/빌드/운영/문제 대응 문서", "비개발자도 APK 설치와 기본 테스트 가능"),
            ("릴리즈", "APK 또는 AAB, 릴리즈 노트, 버전 정책", "semantic version 기반 변경 이력 제공"),
        ],
        (1.1, 2.9, 2.35),
    )

    add_callout(
        doc,
        "개발 범위 요약",
        "Android 앱은 현장 송출과 모바일 관제를 위한 별도 native client입니다. 핵심은 WebRTC 송출/수신, GPS/지도, 이벤트/푸시, 보안 토큰 관리, 실기기 테스트입니다.",
        LIGHT_TEAL,
        TEAL,
    )


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_reference_screens(doc)
    add_requirements_pages(doc)
    add_header_footer(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
    print(PDF_PATH)
