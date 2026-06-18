from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "design"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "GCS-Saker_프론트디자인_업체요구사항_브리프_v0.1.docx"
PDF_PATH = OUT_DIR / "GCS-Saker_프론트디자인_업체요구사항_브리프_v0.1.pdf"

FONT_NAME = "Apple SD Gothic Neo"
FONT_PATH = "/System/Library/Fonts/AppleSDGothicNeo.ttc"

INK = RGBColor(21, 34, 48)
MUTED = RGBColor(82, 96, 112)
BLUE = RGBColor(31, 105, 160)
TEAL = RGBColor(15, 118, 110)
ORANGE = RGBColor(180, 83, 9)
RED = RGBColor(153, 27, 27)
GREEN = RGBColor(22, 101, 52)
LIGHT_BLUE = "EAF3FA"
LIGHT_GRAY = "F4F6F8"
LIGHT_TEAL = "E6F4F1"
LIGHT_ORANGE = "FFF3E7"
LIGHT_RED = "FDECEC"
BORDER = "D3DCE6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(9.5)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color: str = BORDER) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_font(run, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_font(run, color=INK)
    return paragraph


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25 + level * 0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=10.2, color=INK)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=10.2, color=INK)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    size = 16 if level == 1 else 12.5 if level == 2 else 11.2
    color = BLUE if level in (1, 2) else TEAL
    set_font(run, size=size, bold=True, color=color)
    return paragraph


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, color: RGBColor = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_table_borders(table, color="B7C9D8")
    set_cell_margins(cell, top=160, bottom=150, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(title)
    set_font(run, size=10.5, bold=True, color=color)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(body)
    set_font(run, size=9.8, color=INK)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.75, 4.45), header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 if header else 0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    set_table_borders(table)
    if header:
        cells = table.rows[0].cells
        for i, text in enumerate(header):
            set_cell_shading(cells[i], LIGHT_GRAY)
            set_cell_text(cells[i], text, bold=True, color=BLUE)
            set_cell_margins(cells[i])
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], "F8FAFC")
        set_cell_text(cells[0], label, bold=True, color=INK)
        set_cell_text(cells[1], value, color=INK)
        for cell in cells:
            set_cell_margins(cell)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_three_col_table(doc: Document, header: tuple[str, str, str], rows: list[tuple[str, str, str]], widths=(1.5, 2.35, 2.35)):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    set_table_borders(table)
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True, color=BLUE)
        set_cell_margins(cell)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i == 0), color=INK)
            set_cell_margins(cells[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def draw_dashboard_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1800, 1030), "#07131f")
    d = ImageDraw.Draw(img)
    font = ImageFont.truetype(FONT_PATH, 34, index=0)
    small = ImageFont.truetype(FONT_PATH, 24, index=0)
    tiny = ImageFont.truetype(FONT_PATH, 20, index=0)

    def box(x1, y1, x2, y2, title, body, color="#2e87c8", fill="#0c2133"):
        d.rounded_rectangle((x1, y1, x2, y2), radius=18, outline=color, width=4, fill=fill)
        d.text((x1 + 28, y1 + 22), title, font=font, fill="#d9ecff")
        if body:
            lines = body.split("\n")
            for idx, line in enumerate(lines):
                d.text((x1 + 28, y1 + 78 + idx * 31), line, font=small, fill="#aac1d6")

    d.text((48, 36), "기본 관제 대시보드 정보 구조", font=ImageFont.truetype(FONT_PATH, 46, index=0), fill="#ffffff")
    d.line((48, 104, 1752, 104), fill="#2e87c8", width=3)
    box(55, 140, 405, 620, "자산 트리", "조직 / 그룹\n드론 / 로봇 / 센서\n연결 / 끊김 / 경고", "#5bc0de")
    box(430, 140, 1110, 620, "지도", "위성 지도 / 폐쇄망 지도\nGPS 핀 / Auto Focus\n커스텀 마커 / 좌표 표시", "#4ade80", "#0d261f")
    box(1135, 140, 1745, 620, "실시간 스트림", "WebRTC Player Grid\n상태 표시 / AI 모드\nPIN / POP-out / 전체화면", "#38bdf8", "#0b2536")
    box(55, 650, 405, 930, "운영 상태", "API / 인증 / Signaling\nTURN / MediaMTX / DB\n지연 / 연결 수 / 오류율", "#f59e0b", "#2a1d08")
    box(430, 650, 820, 930, "텔레메트리", "선택 스트림 기준\nGPS / 배터리 / 품질\n고도 / 속도 / 업데이트", "#a78bfa", "#211a33")
    box(845, 650, 1220, 930, "이벤트 로그", "검색 / 기간 / 심각도\n서비스 / 자산 / 사용자\n그래프 요약", "#fb7185", "#2a1219")
    box(1245, 650, 1745, 930, "사용자 설정", "닉네임 / 권한\n시간 동기화\n스트림 이름 / 개인 설정", "#2dd4bf", "#092c2a")
    d.text((56, 962), "원칙: 기본 레이아웃은 제공하되, 모든 패널은 추가 / 삭제 / 숨김 / 이동 / 고정 / 확장 가능해야 함", font=tiny, fill="#b8c7d6")
    img.save(path)


def draw_webrtc_flow(path: Path) -> None:
    img = Image.new("RGB", (1800, 900), "#ffffff")
    d = ImageDraw.Draw(img)
    font = ImageFont.truetype(FONT_PATH, 34, index=0)
    small = ImageFont.truetype(FONT_PATH, 23, index=0)
    title = ImageFont.truetype(FONT_PATH, 44, index=0)
    d.text((54, 45), "WebRTC 스트림 연결 상태와 UI 표현", font=title, fill="#152230")

    nodes = [
        (70, 160, 330, 310, "1. 로그인", "세션 확인\n권한 확인", "#EAF3FA"),
        (390, 160, 650, 310, "2. 스트림 선택", "감지 목록\n자산 트리 반영", "#E6F4F1"),
        (710, 160, 970, 310, "3. Signaling", "WHIP/WHEP\nSDP 교환", "#FFF3E7"),
        (1030, 160, 1290, 310, "4. ICE 후보", "STUN 우선\nTURN fallback", "#F4F6F8"),
        (1350, 160, 1610, 310, "5. 재생", "저지연 영상\n오디오 표시", "#EAF3FA"),
    ]
    for x1, y1, x2, y2, head, body, fill in nodes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline="#B7C9D8", width=3)
        d.text((x1 + 22, y1 + 25), head, font=font, fill="#1F69A0")
        for idx, line in enumerate(body.split("\n")):
            d.text((x1 + 24, y1 + 82 + idx * 31), line, font=small, fill="#31445A")
    for x in [338, 658, 978, 1298]:
        d.line((x, 235, x + 42, 235), fill="#1F69A0", width=5)
        d.polygon([(x + 42, 235), (x + 25, 223), (x + 25, 247)], fill="#1F69A0")

    lanes = [
        ("정상", "STUN direct 또는 짧은 TURN 경로로 재생. UI는 녹색 연결 상태와 낮은 지연 표시.", "#166534", "#EAF7EF"),
        ("지연", "ICE 연결은 됐지만 RTT 또는 packet loss가 증가. UI는 경고와 품질 지표를 보여줌.", "#B45309", "#FFF7ED"),
        ("끊김", "publisher 중단, 인증 만료, relay 실패. 재연결 중 / fallback / 로그 이벤트가 표시됨.", "#991B1B", "#FEF2F2"),
    ]
    y = 430
    for label, desc, color, fill in lanes:
        d.rounded_rectangle((85, y, 1625, y + 105), radius=18, fill=fill, outline=color, width=2)
        d.text((115, y + 30), label, font=font, fill=color)
        d.text((245, y + 34), desc, font=small, fill="#263849")
        y += 130
    img.save(path)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("GCS-Saker Frontend Design Brief")
    set_font(run, size=8.5, color=MUTED)
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("A4AI / 디자인 업체 전달용 / v0.1")
    set_font(run, size=8.2, color=MUTED)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    dashboard_png = ASSET_DIR / "gcs_saker_dashboard_information_architecture.png"
    webrtc_png = ASSET_DIR / "gcs_saker_webrtc_ui_state_flow.png"
    draw_dashboard_diagram(dashboard_png)
    draw_webrtc_flow(webrtc_png)

    doc = Document()
    configure_document(doc)
    add_header_footer(doc)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("GCS-Saker")
    set_font(run, size=28, bold=True, color=BLUE)
    p = doc.add_paragraph()
    run = p.add_run("프론트엔드 디자인 업체 요구사항 브리프")
    set_font(run, size=22, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    run = p.add_run("실시간 드론/로봇/센서 관제 대시보드 UI/UX 설계 요청서")
    set_font(run, size=12.5, color=MUTED)
    p.paragraph_format.space_after = Pt(20)

    add_callout(
        doc,
        "한 줄 요약",
        "GCS-Saker는 여러 WebRTC 스트림, 지도, 자산, 서버 상태, 이벤트 로그를 동시에 다루는 저지연 작전 관제 대시보드입니다. "
        "디자인 목표는 화려한 랜딩페이지가 아니라 빠른 판단, 안정적인 운영, 확장 가능한 컴포넌트 구조입니다.",
        LIGHT_BLUE,
        BLUE,
    )

    add_kv_table(
        doc,
        [
            ("문서 목적", "디자인 업체가 범위, 화면, 컴포넌트, 산출물, 제약을 이해하고 바로 작업 계획을 세울 수 있게 하는 발주 브리프"),
            ("대상 사용자", "관제 운영자, 현장 운용자, 관리자, 시스템 운영 담당자"),
            ("주요 환경", "데스크톱 관제 화면, 모바일/노트북 송출 화면, 공개망 및 폐쇄망 납품 환경"),
            ("기술 전제", "React/TypeScript 기반 프론트, WebRTC/HLS, 지도, JWT/httpOnly 세션, 서버 상태 API, 이벤트 로그 API"),
            ("디자인 방향", "다크 관제형 UI, 정보 밀도는 높지만 피로감은 낮게, 상태 변화는 명확하고 조작은 빠르게"),
        ],
    )

    doc.add_page_break()

    add_heading(doc, "1. 제품 맥락과 디자인 원칙")
    add_paragraph(
        doc,
        "GCS-Saker는 드론, 로봇, 센서, 서버 상태를 실시간으로 관제하는 통합 상황판입니다. 사용자는 여러 영상 스트림을 보면서 지도상의 위치, "
        "자산 상태, 네트워크 지연, 서버 오류, 이벤트 로그를 함께 해석해야 합니다. 따라서 첫 화면은 설명용 랜딩페이지가 아니라 실제 작전 관제 화면이어야 합니다.",
    )
    add_three_col_table(
        doc,
        ("원칙", "디자인 의미", "검수 기준"),
        [
            ("빠른 판단", "중요 상태가 1초 안에 식별되어야 함", "정상/경고/장애/재연결 상태가 색상, 아이콘, 텍스트로 동시에 구분됨"),
            ("장시간 사용", "운영자가 오래 봐도 피로가 적어야 함", "과한 애니메이션, 고채도 배경, 과밀한 텍스트를 피함"),
            ("실시간성", "스트림과 상태가 현재성을 가져야 함", "로딩/연결/끊김/fallback 상태가 화면에 즉시 반영됨"),
            ("확장성", "위젯과 자산 종류가 계속 늘어남", "컴포넌트 단위로 재사용되고 grid 배치가 변경 가능함"),
            ("보안 전제", "토큰/비밀값 노출을 전제로 하지 않음", "localStorage 의존, credential 노출 흐름을 요구하지 않음"),
        ],
    )

    add_heading(doc, "2. 기본 화면 정보 구조")
    add_paragraph(
        doc,
        "기본 레이아웃은 기존 관제 시안을 기준으로 하되, 고정 화면이 아니라 사용자가 패널을 추가, 삭제, 숨김, 이동, 고정할 수 있는 grid 기반 구조여야 합니다.",
    )
    doc.add_picture(str(dashboard_png), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        doc,
        "위 구조는 기본값입니다. 사용자가 불필요한 패널을 숨기거나, 스트리밍 영역을 크게 보거나, 이벤트 로그와 서버 상태를 운영 모드로 확장할 수 있어야 합니다.",
    )

    doc.add_page_break()

    add_heading(doc, "3. 핵심 화면 목록")
    add_three_col_table(
        doc,
        ("화면", "핵심 목적", "필수 상태"),
        [
            ("로그인/회원가입", "운영자와 송출자를 명확히 분리하고 세션을 안전하게 시작", "로그인 실패, 세션 만료, 권한 없음, 초대 코드 안내"),
            ("메인 대시보드", "스트림, 지도, 자산, 서버 상태를 동시에 관제", "정상, 경고, 장애, 재연결, fallback"),
            ("스트리밍 선택", "감지된 스트림을 선택하고 재생 위젯에 연결", "감지됨, 연결 중, 연결됨, 끊김, 권한 없음"),
            ("지도 관제", "GPS 기반 자산 위치와 선택 스트림 focus 표시", "Auto focus on/off, 좌표 표시, 핀 선택, 지도 로딩 실패"),
            ("자산 트리", "조직/그룹/장비 계층과 연결 상태 표현", "온라인, 오프라인, 경고, 이동 중, 선택됨"),
            ("이벤트 로그", "운영 이슈를 검색, 필터, 분석", "검색 없음, 필터 적용, 중요 이벤트, 시간 범위"),
            ("사용자/설정", "닉네임, 스트림 이름, 시간 동기화 등 운영 설정", "저장 성공, 실패, 권한 제한"),
        ],
    )

    add_heading(doc, "4. 메인 대시보드 요구사항")
    for item in [
        "첫 화면은 실제 사용 가능한 관제 화면이어야 하며, 제품 설명용 hero 또는 마케팅 화면은 필요하지 않습니다.",
        "기본 배치는 자산 트리, 지도, 스트리밍 grid, 서버 상태, 텔레메트리, 이벤트 로그로 구성합니다.",
        "스트리밍 화면은 실제 영상 판단이 가능한 크기를 확보해야 하며, 지도/텔레메트리가 영상을 압도하면 안 됩니다.",
        "각 패널은 최소 크기를 보장하고, 작은 화면에서도 텍스트와 버튼이 겹치지 않아야 합니다.",
        "위젯 추가, 삭제, 숨김, 초기화, pin, pop-out, 확대, 축소, 순서 이동을 고려합니다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "5. 스트리밍 UI 요구사항")
    add_paragraph(
        doc,
        "스트리밍은 제품의 핵심입니다. 연결 상태가 모호하면 운영자가 시스템 장애와 현장 장애를 구분할 수 없으므로, 모든 상태를 UI 언어로 명확히 표현해야 합니다.",
    )
    doc.add_picture(str(webrtc_png), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_three_col_table(
        doc,
        ("항목", "요구사항", "비고"),
        [
            ("영상 표시", "원본 영상 크기가 아니라 player 영역에 맞춰 표시", "비율 유지, letterbox 허용, 레이아웃 흔들림 금지"),
            ("상태 표시", "연결 중/연결됨/재연결 중/끊김/권한 없음/fallback 표시", "영상 위를 과도하게 가리지 않음"),
            ("오디오", "음성 수신 중이면 스트림 outline 또는 indicator 표시", "깜빡임으로 재렌더링이 과하게 보이면 안 됨"),
            ("조작", "음소거, 전체화면, pin, pop-out, AI 모드 토글", "아이콘 버튼과 tooltip 사용"),
            ("AI 모드", "현재는 예약 기능이지만 overlay/filter가 붙을 수 있게 설계", "켜짐/꺼짐/미지원 상태 필요"),
        ],
    )

    doc.add_page_break()

    add_heading(doc, "6. 지도 UI 요구사항")
    for item in [
        "공개망에서는 위성 지도를 기본으로 사용할 수 있어야 합니다.",
        "폐쇄망 납품을 고려해 외부 지도 API에 완전히 종속되지 않는 구조를 전제로 디자인합니다.",
        "스트림에서 GPS가 들어오면 지도 위에 핀으로 고정 표시해야 합니다.",
        "스트림을 선택하면 기본적으로 해당 위치로 auto focus하고, 사용자가 지도를 직접 움직이면 auto가 꺼집니다.",
        "Auto 버튼을 다시 누르면 선택 스트림의 GPS 위치로 focus가 돌아갑니다.",
        "마우스 위치 좌표 표시, 수동 마커 추가, 마커 이동/고정, 경로 표시 확장이 가능해야 합니다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "7. 자산 트리와 조직/권한 개념")
    add_paragraph(
        doc,
        "자산 트리는 단순 목록이 아니라 조직, 그룹, 장비, 연결 상태를 함께 보여주는 관제 탐색 구조입니다. 향후 군 전술망 또는 기관별 운용을 고려해 상위/하위 조직 개념을 반영할 수 있어야 합니다.",
    )
    add_kv_table(
        doc,
        [
            ("기본 계층", "조직 > 그룹 > 장비 유형 > 개별 장비"),
            ("장비 유형", "드론, 지상 로봇, 센서, 카메라, 서버, 기타 확장 장비"),
            ("상태", "온라인, 오프라인, 연결 중, 경고, 장애, 권한 없음"),
            ("동작", "스트림 감지 시 자동 추가, 끊김 감지 시 상태 변경, 이름 설정, 선택 시 지도/텔레메트리 focus"),
            ("권한 확장", "내 조직, 하위 조직, 공유 자산, 보기 전용, 제어 가능 권한을 구분할 수 있어야 함"),
        ],
    )

    add_heading(doc, "8. 서버 상태와 운영 지표")
    add_three_col_table(
        doc,
        ("서비스", "표시할 정보", "상태 표현"),
        [
            ("API 서버", "health, ready, 응답 시간, 오류율", "정상/지연/장애"),
            ("인증 서버", "로그인 성공률, 토큰 갱신 실패, 세션 오류", "정상/경고/장애"),
            ("Signaling/Media", "WHEP/WHIP 응답, active stream, fallback", "정상/재연결/장애"),
            ("TURN 서버", "primary/secondary 상태, relay 사용 비율, allocation 추정", "정상/부하/장애"),
            ("DB/Redis", "연결 상태, 응답 시간, 장애 degraded 여부", "정상/지연/장애"),
            ("Nginx/Edge", "443 인입, 라우팅, TLS 상태", "정상/경고/장애"),
        ],
    )

    add_heading(doc, "9. 이벤트 로그와 분석 화면")
    for item in [
        "로그는 단순 텍스트 나열이 아니라 운영자가 장애 원인을 좁혀가는 분석 화면이어야 합니다.",
        "필터는 시간 범위, 이벤트 강도, 서비스 종류, 스트림 ID, 사용자/자산, 키워드를 지원해야 합니다.",
        "시간대별 연결 수, 오류 수, 재연결 횟수, 평균 지연, TURN 사용 비율을 그래프 형태로 볼 수 있어야 합니다.",
        "중요 이벤트는 운영자가 놓치지 않되, 전체 관제 흐름을 막는 modal 남발은 피합니다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "10. 디자인 시스템 요구사항")
    add_three_col_table(
        doc,
        ("영역", "요구사항", "산출물"),
        [
            ("Color", "다크 관제형 기본, 경고/장애/정상 상태 색상 체계", "palette, semantic color token"),
            ("Typography", "장시간 관제에 적합한 가독성, 작은 패널 내 과대 제목 금지", "type scale, usage rule"),
            ("Icon", "버튼은 가능한 아이콘 중심, tooltip 제공", "icon set, state icon"),
            ("Component", "stream card, map pin, asset node, status badge, log row", "variant, state, spacing spec"),
            ("Motion", "상태 전환은 부드럽게, 과한 장식 애니메이션 금지", "transition guideline"),
            ("Responsive", "데스크톱 관제 우선, 모바일은 송출/간단 확인 중심", "breakpoint, layout rule"),
        ],
    )

    add_heading(doc, "11. 보안과 기술 제약")
    add_callout(
        doc,
        "중요 제약",
        "JWT, refresh token, TURN credential, 운영자 비밀번호, API key 같은 값이 UI 또는 localStorage에 노출되는 흐름을 전제로 디자인하면 안 됩니다. "
        "세션은 httpOnly cookie 기반으로 이동할 수 있으며, 권한 없음/세션 만료/재로그인 UI가 필요합니다.",
        LIGHT_RED,
        RED,
    )
    for item in [
        "공개망과 폐쇄망 모두 고려합니다. 공개 API 의존 기능은 폐쇄망 대체 상태를 가져야 합니다.",
        "WebRTC 연결은 STUN 우선, TURN fallback을 전제로 하며, relay 사용 중임을 운영자가 이해할 수 있어야 합니다.",
        "영상/음성 명령 데이터, XSS, CSRF, URL 직접 접근을 고려한 인증/인가 흐름이 필요합니다.",
        "실제 구현은 React/TypeScript 컴포넌트로 이어지므로 Figma 컴포넌트 이름과 상태명이 개발 명세와 맞아야 합니다.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "12. 업체 산출물")
    add_three_col_table(
        doc,
        ("산출물", "내용", "검수 기준"),
        [
            ("Figma 원본", "전체 화면, 컴포넌트, variant, prototype", "개발자가 컴포넌트 단위로 추출 가능"),
            ("Design system", "색상, 타이포, 간격, 아이콘, 상태 색상", "semantic token 기준 명확"),
            ("화면 설계", "로그인, 대시보드, 지도, 스트림, 이벤트 로그, 설정", "상태별 화면 포함"),
            ("Interaction spec", "연결/끊김/재연결/fallback/권한 없음 흐름", "운영자가 다음 행동을 알 수 있음"),
            ("Responsive spec", "데스크톱/노트북/태블릿/모바일 송출 화면", "텍스트와 버튼 겹침 없음"),
            ("개발 handoff", "컴포넌트 이름, props 수준 상태, spacing, z-index guide", "React 구현에 바로 연결 가능"),
        ],
    )

    doc.add_page_break()

    add_heading(doc, "13. 검수 체크리스트")
    checklist = [
        "첫 화면이 실제 관제 화면으로 시작하는가",
        "스트리밍 영역이 충분히 크고 영상 위 불필요한 텍스트가 없는가",
        "연결/끊김/재연결/권한 없음/fallback 상태가 명확한가",
        "지도 pin이 GPS 위치에 고정되고 auto focus UX가 이해되는가",
        "자산 트리에서 연결됨과 끊김이 모두 표현되는가",
        "서버 상태가 API/인증/Signaling/TURN/DB/Redis 등으로 분리되는가",
        "이벤트 로그 필터와 그래프가 운영 의사결정에 도움이 되는가",
        "위젯 추가/삭제/숨김/이동/고정 UX가 자연스러운가",
        "보안상 localStorage나 credential 노출을 전제로 하지 않는가",
        "폐쇄망 지도/API 대체 상태가 고려되어 있는가",
    ]
    for item in checklist:
        add_number(doc, item)

    add_heading(doc, "14. 작업 협의 시 확인 질문")
    for item in [
        "Figma 컴포넌트와 개발 컴포넌트 이름을 어느 수준까지 맞출 것인가",
        "관제 화면의 기본 grid 비율과 스트리밍 우선 모드 비율을 어떻게 둘 것인가",
        "장애/경고/정보 알림을 toast, banner, panel 중 어떤 규칙으로 나눌 것인가",
        "위성 지도와 폐쇄망 지도 대체 화면의 시각 차이를 어떻게 최소화할 것인가",
        "모바일 송출 화면을 같은 디자인 시스템 안에서 어디까지 포함할 것인가",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "최종 메시지",
        "이 프로젝트의 좋은 디자인은 예쁜 화면 하나가 아니라, 실제 운영자가 장애와 현장 상황을 빠르게 구분하고 스트림을 안정적으로 다룰 수 있게 만드는 것입니다.",
        LIGHT_TEAL,
        TEAL,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
    print(PDF_PATH)
