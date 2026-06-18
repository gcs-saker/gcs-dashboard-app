from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "design"
DOCX_PATH = OUT_DIR / "GCS-Saker_프론트디자인_업체요구사항_브리프_v0.6.docx"
PDF_PATH = OUT_DIR / "GCS-Saker_프론트디자인_업체요구사항_브리프_v0.6.pdf"
DASHBOARD_SCREENSHOT = OUT_DIR / "assets" / "gcs_saker_current_dashboard_20260608.png"
CCTV_SCREENSHOT = OUT_DIR / "assets" / "gcs_saker_current_cctv_20260608.png"
EVENT_LOG_SCREENSHOT = OUT_DIR / "assets" / "gcs_saker_current_event_log_20260608.png"
SETTINGS_SCREENSHOT = OUT_DIR / "assets" / "gcs_saker_current_settings_20260608.png"

FONT_NAME = "Apple SD Gothic Neo"
INK = RGBColor(22, 32, 44)
MUTED = RGBColor(82, 96, 112)
BLUE = RGBColor(31, 105, 160)
TEAL = RGBColor(15, 118, 110)
RED = RGBColor(153, 27, 27)
LIGHT_BLUE = "EAF3FA"
LIGHT_GRAY = "F5F7FA"
LIGHT_TEAL = "E6F4F1"
LIGHT_RED = "FDECEC"
BORDER = "D3DCE6"


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 9.6) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def configure_section(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_document(doc: Document) -> None:
    configure_section(doc.sections[0], landscape=False)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.13
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("GCS-Saker Frontend Design Brief")
        set_font(run, size=8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("A4AI / 디자인 업체 전달용 / v0.6")
        set_font(run, size=8.2, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else TEAL)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.13
    run = paragraph.add_run(text)
    set_font(run, size=10.4, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=10.1, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, color: RGBColor = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_table_borders(table, color="B7C9D8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=160, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=10.6, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(body)
    set_font(run, size=9.8, color=INK)


def add_table(doc: Document, header: tuple[str, ...], rows: list[tuple[str, ...]], widths: tuple[float, ...]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table)
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, text, bold=True, color=BLUE, size=9.6)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_margins(cells[idx])
            if idx == 0:
                set_cell_shading(cells[idx], "F8FAFC")
            set_cell_text(cells[idx], text, bold=(idx == 0), size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("GCS-Saker")
    set_font(run, size=28, bold=True, color=BLUE)
    p = doc.add_paragraph()
    run = p.add_run("프론트엔드 디자인 브리프")
    set_font(run, size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("현 구현 화면 기반 UI/UX 디자인 제안 요청서")
    set_font(run, size=12.5, color=MUTED)
    add_callout(
        doc,
        "중요: 이 문서는 디자인 작업 범위만 다룹니다",
        "기능 구현, API 연동, 인증/인가, WebRTC 연결, 서버 상태 수집, 지도 데이터 처리 등은 A4AI 개발팀이 담당합니다. "
        "본 문서는 현재 구현 화면을 기준으로 화면 구조, 시각 언어, 컴포넌트 체계, "
        "인터랙션 표현, 개발 handoff 산출물의 디자인 범위를 정리합니다.",
        LIGHT_RED,
        RED,
    )
    add_table(
        doc,
        ("구분", "내용"),
        [
            ("문서 목적", "현 구현 화면과 제품 성격을 기준으로 개선된 디자인 방향과 Figma 산출물 범위 정리"),
            ("작업 범위", "UI/UX, 정보 구조, 디자인 시스템, 컴포넌트 variant, 상태 표현, 프로토타입, handoff 명세"),
            ("작업 제외", "프론트/백엔드 기능 구현, API 개발, WebRTC/지도/서버 연동 구현, 보안 로직 구현"),
            ("중요 관점", "관제 상황에서 빠른 판단, 낮은 피로도, 명확한 상태 인지, 확장 가능한 컴포넌트 설계"),
            ("참고 기준", "대시보드, CCTV, 이벤트로그, 운영설정의 현 구현 화면"),
        ],
        (1.55, 4.75),
    )


def add_screen_reference_page(doc: Document, title: str, body: str, image_path: Path, width: float = 8.6) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    add_header_footer(doc)
    add_heading(doc, title)
    add_paragraph(doc, body)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_callout(
            doc,
            "현 구현 화면 누락",
            f"이미지 파일을 찾을 수 없습니다: {image_path}",
            LIGHT_RED,
            RED,
        )


def add_current_screen_pages(doc: Document) -> None:
    add_screen_reference_page(
        doc,
        "현 구현 화면 1: 메인 대시보드",
        "2026년 6월 8일 배포 사이트에서 캡처한 메인 대시보드입니다. 자산 트리, 지도, 선택 스트림, 다중 스트림, 서버 상태, 텔레메트리, AI 결과 영역의 정보 구조를 확인하기 위한 기준 화면입니다.",
        DASHBOARD_SCREENSHOT,
        width=8.6,
    )
    add_screen_reference_page(
        doc,
        "현 구현 화면 2: CCTV",
        "다중 영상 모니터링 탭입니다. 영상 카드 배치, 빈/오프라인 상태, 선택 상태, 다중 스트림 관제 화면의 밀도와 여백을 검토하기 위한 기준 화면입니다.",
        CCTV_SCREENSHOT,
        width=8.6,
    )
    add_screen_reference_page(
        doc,
        "현 구현 화면 3: 이벤트로그",
        "운영 이벤트와 네트워크/연결 지표를 확인하는 화면입니다. 검색, 필터, 시간 범위, 심각도, 이벤트 목록, 요약 그래프의 정보 구조를 디자인 범위에 포함합니다.",
        EVENT_LOG_SCREENSHOT,
        width=8.6,
    )
    add_screen_reference_page(
        doc,
        "현 구현 화면 4: 운영설정",
        "시간 동기화와 운영 설정을 다루는 화면입니다. 설정 입력, 저장/복구, 상태 확인, 민감 정보가 드러나지 않는 폼 구조와 검수 가능한 상태 표현을 디자인 범위에 포함합니다.",
        SETTINGS_SCREENSHOT,
        width=8.6,
    )


def add_scope_pages(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    add_header_footer(doc)

    add_heading(doc, "1. 디자인 의뢰의 핵심")
    add_paragraph(
        doc,
        "GCS-Saker는 드론, 로봇, 센서, 서버 상태를 동시에 보는 실시간 관제 대시보드입니다. "
        "디자인 목표는 기능을 새로 정의하는 것이 아니라, 현재 기능과 예정 기능이 운영자에게 명확하고 안정적으로 보이도록 화면 체계를 설계하는 것입니다.",
    )
    add_table(
        doc,
        ("디자인 작업 범위", "A4AI 개발팀 구현 범위"),
        [
            ("정보 구조, 화면 우선순위, grid 배치, 시각 비례", "데이터 수집, API, WebRTC signaling, 서버 로직"),
            ("상태 표현: 정상, 지연, 경고, 장애, 연결/끊김", "실제 상태 판정 로직과 이벤트 생성"),
            ("컴포넌트 스타일과 Figma variant", "React/TypeScript 구현과 테스트"),
            ("지도/스트림/자산/로그의 시각적 사용성", "지도 SDK 연동, GPS 처리, 스트림 라우팅"),
            ("디자인 시스템과 개발 handoff", "보안, 인증/인가, 배포, 성능 최적화"),
        ],
        (3.05, 3.05),
    )

    add_heading(doc, "2. 디자인 제안 방향")
    for item in [
        "현 구현 화면의 장단점을 기준으로 보수적 개선안과 더 과감한 재배치안을 비교 제안해 주세요.",
        "다크 관제형 톤은 유지하되, 단조로운 색상이나 과한 장식이 아니라 상태 인지와 장시간 사용성을 우선해 주세요.",
        "스트리밍 영역은 실제 판단이 가능한 크기와 비례를 확보하고, 지도/로그/상태 패널은 스트림을 방해하지 않게 구성해 주세요.",
        "사용자가 패널을 추가, 숨김, 고정, 확장, pop-out 할 수 있는 grid 기반 조작감을 디자인해 주세요.",
        "데스크톱 관제 화면을 우선 설계하되, 모바일/노트북 송출 화면의 디자인 방향도 같은 시스템 안에서 제안해 주세요.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    add_heading(doc, "3. 화면별 디자인 대상")
    add_table(
        doc,
        ("영역", "디자인 과제", "주의점"),
        [
            ("메인 대시보드", "전체 정보 밀도, 패널 비례, 시선 흐름, 위젯 조작", "관제 사용성 기준으로 스트림/지도/상태 우선순위 개선"),
            ("CCTV", "다중 영상 카드, 빈 상태, 선택/확대, 영상 카드 밀도", "스트림이 없는 상태도 운영 화면처럼 자연스러워야 함"),
            ("스트리밍 패널", "영상 영역 크기, 상태 badge, 오디오 표시, AI 모드 버튼, pin/pop/fullscreen 버튼", "텍스트가 영상 위를 과도하게 가리면 안 됨"),
            ("지도 패널", "위성 지도 톤, 핀/마커/좌표 표시, auto focus 버튼, 선택 상태", "지도 요소가 다른 선택창보다 위로 튀지 않게 z-index 질서 필요"),
            ("자산 트리", "조직/그룹/장비 계층, 연결/끊김/경고 상태, 선택 상태", "복잡한 계층도 빠르게 접고 펼칠 수 있어야 함"),
            ("서버 상태", "API, 인증, Signaling, Media, TURN, DB/Redis 상태 카드", "단순 장식 수치가 아니라 운영 상태가 읽혀야 함"),
            ("이벤트 로그", "검색, 시간 범위, 심각도, 서비스 필터, 그래프 요약", "로그 나열이 아니라 장애 분석 화면처럼 보여야 함"),
            ("운영설정", "사용자 정보, 스트림 이름, 시간 동기화, 개인 설정", "보안상 토큰/비밀값 노출을 전제로 하지 않음"),
        ],
        (1.35, 3.0, 2.0),
    )

    doc.add_page_break()

    add_heading(doc, "4. 디자인 시스템 산출물")
    add_table(
        doc,
        ("산출물", "요구 내용", "검수 기준"),
        [
            ("Figma 원본", "전체 화면, component, variant, prototype 포함", "개발팀이 컴포넌트 단위로 바로 참조 가능"),
            ("Design system", "색상, 타이포, 간격, icon, elevation, border, motion", "semantic token과 상태 색상 기준 명확"),
            ("Component library", "stream card, asset node, map marker, status badge, log row, toolbar", "normal/hover/active/disabled/loading/error 상태 포함"),
            ("Layout spec", "기본 grid, 최소 크기, 확장/축소, pop-out, pinned 상태", "화면 크기별 겹침 없이 동작하도록 정의"),
            ("Interaction spec", "클릭, 선택, hover, focus, panel resize, modal/drawer/toast", "운영자가 다음 행동을 예측할 수 있음"),
            ("Handoff note", "컴포넌트 명명, spacing, z-index, 상태명, responsive rule", "React 구현자와 디자이너가 같은 용어 사용"),
        ],
        (1.55, 3.0, 2.05),
    )

    add_heading(doc, "5. 디자인 검수 기준")
    for item in [
        "현 구현 화면보다 스트리밍 영역의 판단 가능성이 좋아졌는가",
        "지도, 자산, 서버 상태, 로그가 서로 겹치지 않고 우선순위가 분명한가",
        "정상/지연/경고/장애/연결/끊김 상태가 색상만이 아니라 형태와 텍스트로도 구분되는가",
        "사용자가 패널을 커스텀할 수 있다는 조작감이 명확한가",
        "Figma 컴포넌트가 실제 React 컴포넌트로 분리되기 쉬운 구조인가",
        "보안상 localStorage, token, credential 노출을 전제로 한 화면 흐름이 없는가",
        "공개망/폐쇄망 환경 차이가 디자인 상태로 표현 가능한가",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "디자인 방향 요약",
        "현 구현 화면을 바탕으로 GCS-Saker의 관제 경험을 정교화합니다. "
        "핵심 범위는 화면 구조, 시각 언어, 컴포넌트 시스템, 상태 표현, 인터랙션 프로토타입, 개발 handoff입니다.",
        LIGHT_TEAL,
        TEAL,
    )


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_current_screen_pages(doc)
    add_scope_pages(doc)
    add_header_footer(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
    print(PDF_PATH)
