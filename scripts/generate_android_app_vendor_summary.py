from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "design"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "GCS-Saker_안드로이드앱_업체전달용_요약요구사항서_v0.4.docx"
PDF_PATH = OUT_DIR / "GCS-Saker_안드로이드앱_업체전달용_요약요구사항서_v0.4.pdf"

SCREENSHOTS = [
    ("메인 대시보드", ASSET_DIR / "gcs_saker_current_dashboard_20260608.png"),
    ("CCTV", ASSET_DIR / "gcs_saker_current_cctv_20260608.png"),
    ("이벤트로그", ASSET_DIR / "gcs_saker_current_event_log_20260608.png"),
    ("운영설정", ASSET_DIR / "gcs_saker_current_settings_20260608.png"),
]

FONT_NAME = "Apple SD Gothic Neo"
INK = RGBColor(22, 32, 44)
MUTED = RGBColor(82, 96, 112)
BLUE = RGBColor(31, 105, 160)
TEAL = RGBColor(15, 118, 110)
RED = RGBColor(153, 27, 27)
LIGHT_BLUE = "EAF3FA"
LIGHT_RED = "FDECEC"
BORDER = "D3DCE6"


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def configure_section(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def configure_document(doc: Document) -> None:
    configure_section(doc.sections[0])
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.14
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(3)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("GCS-Saker Android App Vendor Brief")
        set_font(run, size=8.5, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("A4AI / 업체 전달용 / v0.4")
        set_font(run, size=8.2, color=MUTED)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: float = 9.3) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color or INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    set_font(run, size=16 if level == 1 else 12.5, bold=True, color=BLUE if level == 1 else TEAL)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.14
    run = paragraph.add_run(text)
    set_font(run, size=10.4, color=INK)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=10.1, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, color: RGBColor = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.35)
    set_table_borders(table, color="B7C9D8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_font(run, size=10.6, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(body)
    set_font(run, size=9.8, color=INK)


def add_table(doc: Document, header: tuple[str, ...], rows: list[tuple[str, ...]], widths: tuple[float, ...]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    set_table_borders(table)
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_text(cell, text, bold=True, color=BLUE, size=9.5)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_margins(cells[idx])
            if idx == 0:
                set_cell_shading(cells[idx], "F8FAFC")
            set_cell_text(cells[idx], text, bold=(idx == 0), size=9.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("GCS-Saker")
    set_font(run, size=28, bold=True, color=BLUE)
    p = doc.add_paragraph()
    run = p.add_run("Android 모바일 앱 개발 의뢰 요약")
    set_font(run, size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("업체 견적 및 제안 요청을 위한 전달용 문서")
    set_font(run, size=12.5, color=MUTED)
    add_callout(
        doc,
        "의뢰 목적",
        "현재 보유 중인 GCS-Saker 웹사이트의 모바일 버전을 Android 앱 형태로 제작합니다. 웹 화면의 기능과 사용 흐름을 모바일 환경에 맞게 재구성하고, 설치 가능한 APK를 산출물로 납품합니다.",
        LIGHT_RED,
        RED,
    )
    add_table(
        doc,
        ("구분", "전달 내용"),
        [
            ("개발 범위", "현재 웹사이트의 모바일 앱 버전 기획 보완, UI/UX 설계, 앱 개발, 서버 연동, 테스트, APK 납품"),
            ("주요 사용자", "현장 송출자, 모바일 관제 사용자, 운영 관리자"),
            ("주요 기능", "로그인, 카메라 송출, 스트림 수신, 지도/GPS, 이벤트/알림, 운영 상태 확인"),
            ("제외 범위", "웹 대시보드 개선, 백엔드 신규 구축, 서버 인프라 운영, iOS 앱, Figma 별도 산출물"),
            ("협업 방식", "A4AI가 서버/API/테스트 계정/연동 환경을 제공하고 업체가 앱 구현과 검증을 담당"),
        ],
        (1.45, 4.85),
    )


def add_scope_page(doc: Document) -> None:
    doc.add_page_break()
    add_header_footer(doc)
    add_heading(doc, "1. 필수 요구사항")
    add_table(
        doc,
        ("구분", "요구사항", "A4AI 제공물"),
        [
            ("모바일 앱화", "현재 웹사이트의 핵심 기능을 Android 모바일 앱 화면 흐름으로 재구성", "현 웹사이트 URL, 검토 계정, 주요 화면 캡처"),
            ("계정/권한", "로그인, 로그아웃, 세션 유지, 역할/그룹 권한 반영", "인증 API endpoint, 테스트 계정, 권한 정책 설명"),
            ("모바일 관제", "스트림 목록, 선택 스트림 재생, 지도 위치, 이벤트 확인", "스트림 목록 API, 이벤트 API, 지도/위치 데이터 형식"),
            ("현장 송출", "휴대폰 카메라/마이크/GPS를 이용해 서버로 실시간 송출", "송출 endpoint, WebRTC signaling 절차, 테스트 스트림 환경"),
            ("Signaling/ICE", "WebRTC 연결 시 GCS-Saker 서버가 제공하는 ICE 서버 목록을 사용", "STUN/TURN 서버 정보, ICE 목록 API 또는 설정값, signaling 예시 코드"),
            ("폐쇄망 알림", "Firebase 등 인터넷용 외부 API를 지양하고 폐쇄망 구현 가능 알림 구조 제안", "알림 이벤트 API, WebSocket 또는 내부 알림 연동 방식 협의"),
            ("디자인 일관성", "현재 웹사이트와 디자인 톤은 유지하되 모바일 화면에 맞게 재배치", "웹 화면 캡처, 색상/컴포넌트 기준, 운영 사용 시나리오"),
            ("APK 납품", "검수 가능한 APK와 소스코드, 테스트 결과, 운영 문서 납품", "검수 시나리오, 테스트 계정, 서버 접속 환경"),
        ],
        (1.1, 2.65, 2.55),
    )

    doc.add_page_break()

    add_heading(doc, "2. 추가 요구사항")
    add_table(
        doc,
        ("구분", "추가 요구사항", "A4AI 제공물"),
        [
            ("운영 상태", "서버 연결 상태, 시간 동기화 상태, 앱 진단 정보 확인 화면", "운영 상태 API, 시간 동기화 API"),
            ("고급 알림", "장애 등급별 알림, 알림 이력, 읽음 처리, 재알림 정책", "이벤트 등급 정책, 알림 메시지 샘플"),
            ("오프라인 대응", "일시적 네트워크 단절 시 상태 표시와 재연결 안내", "장애/복구 테스트 조건"),
            ("스트림 품질", "네트워크 상태에 따른 품질 표시 또는 품질 선택", "스트림 품질 지표, 테스트 스트림"),
            ("관리 기능", "사용자 설정, 서버 선택, 진단 로그 내보내기", "설정 항목 정의, 로그 보안 기준"),
        ],
        (1.15, 2.65, 2.5),
    )

    add_heading(doc, "3. 업체가 제안서에 포함해야 할 것")
    add_table(
        doc,
        ("항목", "확인 내용"),
        [
            ("개발 방식", "앱 구현 방식, 화면 설계 방식, 서버 연동 방식"),
            ("일정", "기획/디자인/개발/테스트/검수 단계별 일정"),
            ("인력", "PM, Android 개발자, UI/UX, QA 역할"),
            ("테스트", "실기기 테스트 범위, WebRTC 송출/수신 검증 방식"),
            ("폐쇄망 대응", "외부 인터넷 API 없이 알림과 연결 상태를 처리하는 방식"),
            ("STUN/TURN", "A4AI STUN/TURN 서버와 ICE 설정을 사용하는 방식"),
            ("산출물", "소스코드, APK, 테스트 결과, 운영 문서"),
            ("유지보수", "하자 보수 기간, 장애 대응 방식, 추가 개발 단가"),
        ],
        (1.35, 4.95),
    )


def add_reference_pair_page(
    doc: Document,
    title: str,
    body: str,
    left: tuple[str, Path],
    right: tuple[str, Path],
) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    add_header_footer(doc)
    add_heading(doc, title)
    add_paragraph(doc, body)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="EEF2F6")
    for idx, (caption, path) in enumerate((left, right)):
        cell = table.cell(0, idx)
        set_cell_margins(cell, top=80, bottom=80, start=90, end=90)
        p = cell.paragraphs[0]
        run = p.add_run(caption)
        set_font(run, size=9.2, bold=True, color=BLUE)
        if path.exists():
            cell.add_paragraph().add_run().add_picture(str(path), width=Inches(4.3))


def add_screen_requirements_page(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    add_header_footer(doc)
    add_heading(doc, "4. 화면별 요구사항")
    add_paragraph(
        doc,
        "아래 항목은 모바일 앱에서 최소로 구성되어야 하는 화면 단위 요구사항입니다. 업체는 동일 기능을 더 나은 모바일 흐름으로 재구성할 수 있으나, 각 화면의 목적과 핵심 상태는 유지해야 합니다.",
    )
    add_table(
        doc,
        ("화면", "필수 구성", "상태/동작", "A4AI 제공물"),
        [
            ("로그인", "아이디/비밀번호 입력, 로그인 유지, 로그아웃 진입", "실패 사유, 세션 만료, 권한 부족 표시", "인증 API, 테스트 계정, 권한 정책"),
            ("홈/대시보드", "서버 상태, 연결 상태, 주요 스트림, 최근 이벤트 요약", "정상/지연/장애 상태가 한눈에 보여야 함", "운영 상태 API, 이벤트 API, 화면 캡처"),
            ("스트림 목록/재생", "스트림 목록, 선택 스트림 재생, 오디오 상태, 연결 상태", "오프라인, 연결중, 재연결, 오류 상태 표시", "스트림 목록 API, WHEP/signaling 정보"),
            ("현장 송출", "카메라/마이크/GPS 권한, 송출 시작/중지, 송출 상태", "signaling 준비, ICE 연결, 송출중, 끊김, 재시도 표시", "WHIP/signaling 절차, STUN/TURN 설정, 테스트 서버"),
            ("지도/GPS", "현재 위치, 스트림 위치, 선택 스트림 포커스", "위치 권한 거부, GPS 수신 실패, 지도 로딩 실패 처리", "위치 데이터 형식, 지도 사용 정책"),
            ("이벤트/알림", "이벤트 목록, 심각도, 시간, 상세 보기, 알림 확인", "폐쇄망에서도 가능한 알림 구조 우선", "이벤트 API, 알림 메시지 샘플, WebSocket 협의"),
            ("운영설정/진단", "서버 주소, 연결 상태, 시간 동기화, 앱 버전, 진단 로그", "비밀값 노출 없이 점검 가능한 형태", "운영 설정 API, 시간 API, 로그 보안 기준"),
        ],
        (1.15, 2.55, 2.55, 2.25),
    )


def add_reference_page(doc: Document) -> None:
    add_reference_pair_page(
        doc,
        "참고 화면 1",
        "현재 GCS-Saker 웹 관제 시스템입니다. Android 앱은 이 화면과 디자인 일관성을 유지하되, 모바일 사용성에 맞게 별도 설계합니다.",
        SCREENSHOTS[0],
        SCREENSHOTS[1],
    )
    add_reference_pair_page(
        doc,
        "참고 화면 2",
        "이벤트와 운영 설정도 모바일 앱에서 접근 가능한 구조로 재해석합니다.",
        SCREENSHOTS[2],
        SCREENSHOTS[3],
    )


def add_delivery_page(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    add_header_footer(doc)
    add_heading(doc, "5. 납품 기준")
    add_table(
        doc,
        ("산출물", "기준"),
        [
            ("앱 소스코드", "Android Studio에서 빌드 가능한 전체 프로젝트"),
            ("설치 파일", "테스트 및 검수 가능한 APK"),
            ("화면 산출물", "주요 화면 구성, 화면 흐름, 디자인 적용 결과"),
            ("테스트 결과", "실기기 테스트, 로그인, 송출, 수신, 알림, 위치 연동 결과"),
            ("운영 문서", "설치 방법, 환경 설정, 장애 대응, 릴리즈 노트"),
        ],
        (1.55, 4.75),
    )

    add_heading(doc, "6. 중요 조건")
    for item in [
        "비밀번호, 토큰, 서버 비밀값은 앱 코드와 로그에 노출하지 않는다.",
        "카메라, 마이크, 위치 권한은 사용자에게 명확히 안내한다.",
        "Firebase 등 인터넷 연결을 전제로 하는 외부 API 사용은 지양하고, 폐쇄망에서도 구현 가능한 알림 구조를 우선 제안한다.",
        "현재 웹사이트와 디자인 일관성은 유지하되, Figma 별도 산출물은 필수로 요구하지 않는다.",
        "스트리밍 연결 실패, 네트워크 끊김, 권한 거부 상황을 화면에서 이해할 수 있게 표시한다.",
        "실제 Android 기기에서 송출과 수신이 되는 것을 검수 기준에 포함한다.",
        "추후 폐쇄망 환경과 자체 STUN/TURN 사용 가능성을 고려한 구조로 설계한다.",
    ]:
        add_bullet(doc, item)

    add_callout(
        doc,
        "한 줄 요약",
        "현재 GCS-Saker 웹사이트의 모바일 앱 버전을 Android APK 형태로 개발하는 건이며, 핵심은 모바일 관제, 현장 송출, 폐쇄망 대응 알림, 보안, 실기기 검증입니다.",
        LIGHT_BLUE,
        BLUE,
    )


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_scope_page(doc)
    add_screen_requirements_page(doc)
    add_reference_page(doc)
    add_delivery_page(doc)
    add_header_footer(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
    print(PDF_PATH)
